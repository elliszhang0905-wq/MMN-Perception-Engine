import io
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

import server
from creator_script_generation import PLATFORM_RULES, draft_prompt, export_script_docx


ASSET = {
    "id": "creator-dna-1",
    "account_name": "超哥超车",
    "platform": "抖音",
    "sample_count": 120,
    "account_positioning": "从真实购车矛盾切入，给出理性判断与验证动作。",
    "topic_formula": "场景矛盾 -> 车型判断 -> 证据边界 -> 行动建议",
    "script_template": {"opening": "用户问题开场", "body": "观点拆解", "ending": "验证动作"},
    "language_rules": ["说人话", "保留判断边界"],
    "transfer_boundary": "只迁移方法论，不复制原文或个人身份。",
}

EVIDENCE = [{
    "id": "evidence-1",
    "title": "家庭长途选车要先看补能和后排体验",
    "coreTopic": "家庭场景",
    "argumentStructure": "问题到证据再到边界",
    "sourceUrl": "https://example.com/evidence-1",
}]


def fake_model(stage, messages):
    if stage == "draft":
        return {
            "title": "一家人跑长途，选车别只盯着续航数字",
            "openingHook": "一家四口跑一趟长途，真正让人累的可能不是里程，而是每次补能和后排的状态。",
            "spokenScript": "这是初稿。",
            "subtitleHighlights": ["长途体验不是一张续航表"],
            "visualSuggestions": [{"timing": "0-3秒", "shot": "一家人装行李", "subtitle": "长途选车先问这个"}],
            "evidenceBoundaries": ["具体配置以官方在售版本为准"],
        }
    if stage == "review":
        return {"verdict": "pass", "issues": [], "factualRisks": [], "revisionInstructions": [], "humanToneChecks": ["口语自然"]}
    return {
        "title": "一家人跑长途，选车别只盯着续航数字",
        "openingHook": "一家四口跑一趟长途，真正让人累的可能不是里程，而是每次补能和后排的状态。",
        "spokenScript": (
            "周末一家四口出门，后备厢塞满行李，孩子在后排睡着，这时候你会发现，长途体验根本不是一张续航表。"
            "先看补能是不是顺路、等待时间能不能接受，再看后排坐姿、空调和高速噪音。车型参数只能回答它有什么，"
            "真实场景才会告诉你这些东西是不是刚好适合你的家庭。试车时别只绕店一圈，带上常用的儿童座椅和行李，"
            "把二排、后备厢和高速路段都走一遍。具体配置和权益会变化，下单前再按官方在售版本核对。"
        ),
        "subtitleHighlights": ["长途体验不是一张续航表", "补能路径比单一数字更重要", "带着真实装备去试车", "下单前核对在售配置"],
        "visualSuggestions": [
            {"timing": "0-3秒", "shot": "一家人把行李装进后备厢", "subtitle": "长途选车别只看续航"},
            {"timing": "4-18秒", "shot": "后排儿童座椅和空调出风口特写", "subtitle": "看真实家庭场景"},
            {"timing": "19-45秒", "shot": "导航补能路线与服务区画面", "subtitle": "补能是否顺路"},
            {"timing": "46-70秒", "shot": "试驾清单逐项打勾", "subtitle": "用自己的装备验证"},
        ],
        "evidenceBoundaries": ["不使用未提供的实测数据", "具体配置和权益以官方在售版本为准"],
        "qualityNote": "已完成平台适配、事实边界与自然表达复核。",
    }


class CreatorScriptGenerationTest(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.db_patch = patch.object(server, "DB_PATH", Path(self.temp.name) / "creator-script.db")
        self.db_patch.start()
        server.init_db()

    def tearDown(self):
        self.db_patch.stop()
        self.temp.cleanup()

    def create_job(self, **overrides):
        body = {
            "edition": "china",
            "creatorAssetId": ASSET["id"],
            "platform": "douyin",
            "brand": "上汽奥迪",
            "model": "奥迪E7X",
            "focus": "家庭长途场景中的补能与后排体验",
            "title": "家庭长途怎么选车",
            **overrides,
        }
        with patch.object(server, "creator_script_asset_context", return_value=(ASSET, EVIDENCE)):
            return server.create_creator_script_job(body, org_id="org-1", start_worker=False)

    def test_all_platforms_have_independent_generation_rules(self):
        self.assertEqual(set(PLATFORM_RULES), {"douyin", "wechat_channels", "bilibili", "xiaohongshu"})
        for platform, rule in PLATFORM_RULES.items():
            prompt = draft_prompt({"platform": platform, "brand": "上汽奥迪", "model": "奥迪E7X", "focus": "家庭长途"}, ASSET, EVIDENCE)
            text = prompt[-1]["content"]
            self.assertIn(rule["label"], text)
            self.assertIn(rule["playbook"], text)
            self.assertIn("禁用空泛开场", text)
            self.assertIn("不是产品事实证据", prompt[0]["content"])
            self.assertIn("只能改写成到店或拍摄时的验证动作", prompt[0]["content"])

    def test_job_persists_progress_result_and_revision(self):
        job = self.create_job()
        completed = server.run_creator_script_job(job["id"], runner=fake_model)
        self.assertEqual(completed["status"], "completed")
        self.assertEqual(completed["progress"], 100)
        self.assertEqual([item["stage"] for item in completed["stages"]], ["brief", "brief", "draft", "review", "final", "delivery"])
        self.assertIn("完成", completed["result"]["qualityNote"])
        latest = server.latest_creator_script_job(ASSET["id"], edition="china", org_id="org-1")
        self.assertEqual(latest["id"], job["id"])
        revised = server.create_creator_script_job(
            {"revisionRequest": "开头更直接，增加家庭长途的具体动作", "platform": "xiaohongshu"},
            org_id="org-1", start_worker=False, parent_job=completed,
        )
        self.assertEqual(revised["revisionNo"], 2)
        self.assertEqual(revised["parentJobId"], job["id"])
        self.assertEqual(revised["request"]["revisionRequest"], "开头更直接，增加家庭长途的具体动作")
        self.assertEqual(revised["platformLabel"], "小红书")

    def test_failure_reason_is_public_safe_and_retryable(self):
        job = self.create_job()

        def fail(stage, messages):
            raise ValueError("未配置 DASHSCOPE_API_KEY，Qwen unavailable")

        failed = server.run_creator_script_job(job["id"], runner=fail)
        self.assertEqual(failed["status"], "failed")
        self.assertIn("生成能力暂不可用", failed["error"])
        self.assertNotRegex(failed["error"], r"Qwen|DASHSCOPE|API_KEY")
        with patch.object(server, "Thread") as thread:
            retried = server.retry_creator_script_job(job["id"], org_id="org-1", runner=fake_model)
        self.assertEqual(retried["status"], "queued")
        self.assertEqual(retried["progress"], 0)
        thread.assert_called_once()

    def test_completed_script_exports_a_real_word_document(self):
        job = self.create_job()
        completed = server.run_creator_script_job(job["id"], runner=fake_model)
        data = export_script_docx(completed)
        self.assertTrue(data.startswith(b"PK"))
        document = Document(io.BytesIO(data))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("一家人跑长途", text)
        self.assertIn("完整口播稿", text)
        self.assertIn("字幕重点", text)
        self.assertEqual(len(document.tables), 1)
        self.assertFalse(document.core_properties.author)

    def test_frontend_contains_complete_workbench_controls(self):
        app = (Path(__file__).resolve().parents[1] / "app.js").read_text(encoding="utf-8")
        css = (Path(__file__).resolve().parents[1] / "style.css").read_text(encoding="utf-8")
        for phrase in ("选择达人", "抖音", "视频号", "B站", "小红书", "完整口播稿", "字幕重点", "画面建议", "一键复制", "导出Word", "修改要求", "失败原因"):
            self.assertIn(phrase, app)
        self.assertIn("creator-script-dialog", css)
        self.assertIn("script-jobs", app)


if __name__ == "__main__":
    unittest.main()
