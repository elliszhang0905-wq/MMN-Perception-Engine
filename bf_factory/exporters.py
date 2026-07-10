"""BF Word导出。使用 compact_reference_guide + customer_pack 首屏结构。"""

import io
import os
import platform
import zipfile
from datetime import datetime
from xml.sax.saxutils import escape

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    HAS_PYTHON_DOCX = True
except ImportError:
    Document = None
    WD_ALIGN_PARAGRAPH = OxmlElement = qn = Inches = Pt = RGBColor = None
    HAS_PYTHON_DOCX = False


ACCENT = RGBColor(20, 91, 110) if HAS_PYTHON_DOCX else "145B6E"
INK = RGBColor(26, 34, 39) if HAS_PYTHON_DOCX else "1A2227"
MUTED = RGBColor(90, 103, 110) if HAS_PYTHON_DOCX else "5A676E"
DOCX_FONT = os.getenv(
    "MMN_BF_DOCX_CJK_FONT",
    "Hiragino Sans GB" if platform.system() == "Darwin" else "Noto Sans CJK SC",
)


def export_brief_docx(*, payload, sections, include_internal=False):
    if not HAS_PYTHON_DOCX:
        return _export_minimal_ooxml(payload, sections, include_internal)
    document = Document()
    _configure_document(document)
    _add_customer_pack_header(document, payload)
    for section in sections:
        if section.get("visibility") == "INTERNAL" and not include_internal:
            continue
        heading = document.add_paragraph(style="Heading 1")
        heading.add_run(section.get("title") or section.get("intent") or "BF章节")
        _add_markdown_body(document, section.get("body") or "")
    _set_footer(document)
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.comments = ""
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    _set_style_font(normal, DOCX_FONT, DOCX_FONT, 11, INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, ACCENT),
        ("Heading 2", 13, 14, 7, ACCENT),
        ("Heading 3", 12, 10, 5, RGBColor(31, 77, 120)),
    ):
        style = document.styles[style_name]
        _set_style_font(style, DOCX_FONT, DOCX_FONT, size, color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        _set_style_font(style, DOCX_FONT, DOCX_FONT, 11, INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def _add_customer_pack_header(document, payload):
    strategy = payload.get("strategy") or {}
    classification = payload.get("classification") or {}
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("MMN BF FACTORY · BRAND COMMERCIAL CONTENT BRIEF")
    _set_run_font(run, DOCX_FONT, DOCX_FONT, 9, ACCENT, True)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run(strategy.get("bfName") or "品牌商业化内容BF")
    _set_run_font(run, DOCX_FONT, DOCX_FONT, 28, INK, True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run(
        f"{strategy.get('brand') or '品牌待确认'} · {strategy.get('model') or '车型待确认'} · "
        f"{classification.get('bfTypeLabel') or classification.get('bfType') or '自适应BF'}"
    )
    _set_run_font(run, DOCX_FONT, DOCX_FONT, 12, MUTED, False)

    metadata = document.add_table(rows=2, cols=2)
    metadata.autofit = False
    widths = [Inches(3.25), Inches(3.25)]
    values = [
        ("项目阶段", strategy.get("projectStage") or "待确认"),
        ("传播目标", "、".join(strategy.get("communicationGoals") or []) or "待确认"),
        ("核心竞品", "、".join(strategy.get("competitors") or []) or "待确认"),
        ("版本日期", datetime.now().strftime("%Y-%m-%d")),
    ]
    for index, (label, value) in enumerate(values):
        row, col = divmod(index, 2)
        cell = metadata.cell(row, col)
        cell.width = widths[col]
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        label_run = paragraph.add_run(f"{label}：")
        _set_run_font(label_run, DOCX_FONT, DOCX_FONT, 9.5, MUTED, True)
        value_run = paragraph.add_run(str(value))
        _set_run_font(value_run, DOCX_FONT, DOCX_FONT, 9.5, INK, False)
        _set_cell_margins(cell, 80, 120, 80, 120)
    _set_table_geometry(metadata, [4680, 4680], 9360, 120)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def _add_markdown_body(document, body):
    for raw_line in str(body or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith(">"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(line.lstrip("> "))
            _set_run_font(run, DOCX_FONT, DOCX_FONT, 10, MUTED, False, italic=True)
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(line[2:].strip())
        elif line[:2].rstrip(".").isdigit() and ". " in line[:4]:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(line.split(". ", 1)[1])
        else:
            paragraph = document.add_paragraph()
            paragraph.add_run(line)


def _set_footer(document):
    section = document.sections[0]
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("MMN BF工厂 · 品牌商业化内容Brief · ")
    _set_run_font(run, DOCX_FONT, DOCX_FONT, 8.5, MUTED, False)
    _append_page_field(paragraph)


def _append_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def _set_style_font(style, ascii_name, east_asia_name, size, color):
    style.font.name = ascii_name
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style._element.rPr.rFonts.set(qn("w:ascii"), ascii_name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_name)
    style._element.rPr.rFonts.set(qn("w:cs"), east_asia_name)
    style._element.rPr.rFonts.set(qn("w:hint"), "eastAsia")


def _set_run_font(run, ascii_name, east_asia_name, size, color, bold, italic=False):
    run.font.name = ascii_name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), east_asia_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hint"), "eastAsia")


def _set_cell_margins(cell, top, start, bottom, end):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths, total_width, indent):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), str(total_width))
    tbl_width.set(qn("w:type"), "dxa")
    tbl_indent = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_indent)
    tbl_indent.set(qn("w:w"), str(indent))
    tbl_indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:w"), str(widths[index]))
            tc_width.set(qn("w:type"), "dxa")


def _export_minimal_ooxml(payload, sections, include_internal):
    strategy = payload.get("strategy") or {}
    classification = payload.get("classification") or {}
    body = []
    body.append(_ooxml_paragraph("MMN BF FACTORY · BRAND COMMERCIAL CONTENT BRIEF", style="Kicker"))
    body.append(_ooxml_paragraph(strategy.get("bfName") or "品牌商业化内容BF", style="DocTitle"))
    subtitle = (
        f"{strategy.get('brand') or '品牌待确认'} · {strategy.get('model') or '车型待确认'} · "
        f"{classification.get('bfTypeLabel') or classification.get('bfType') or '自适应BF'}"
    )
    body.append(_ooxml_paragraph(subtitle, style="Subtitle"))
    body.append(_ooxml_paragraph(f"项目阶段：{strategy.get('projectStage') or '待确认'}"))
    body.append(_ooxml_paragraph(f"传播目标：{'、'.join(strategy.get('communicationGoals') or []) or '待确认'}"))
    body.append(_ooxml_paragraph(f"核心竞品：{'、'.join(strategy.get('competitors') or []) or '待确认'}"))
    for section in sections:
        if section.get("visibility") == "INTERNAL" and not include_internal:
            continue
        body.append(_ooxml_paragraph(section.get("title") or section.get("intent") or "BF章节", style="Heading1"))
        for raw_line in str(section.get("body") or "").splitlines():
            line = raw_line.strip()
            if not line:
                continue
            if line.startswith("- "):
                body.append(_ooxml_paragraph(line[2:], numbering=(1, 0)))
            elif re_numbered(line):
                body.append(_ooxml_paragraph(line.split(". ", 1)[1], numbering=(2, 0)))
            elif line.startswith(">"):
                body.append(_ooxml_paragraph(line.lstrip("> "), style="Quote"))
            else:
                body.append(_ooxml_paragraph(line))
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:body>' + "".join(body) +
        '<w:sectPr><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="708" w:footer="708"/></w:sectPr>'
        '</w:body></w:document>'
    )
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", _CONTENT_TYPES_XML)
        archive.writestr("_rels/.rels", _ROOT_RELS_XML)
        archive.writestr("docProps/core.xml", _CORE_XML)
        archive.writestr("docProps/app.xml", _APP_XML)
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/styles.xml", _STYLES_XML.replace("{{DOCX_FONT}}", escape(DOCX_FONT)))
        archive.writestr("word/numbering.xml", _NUMBERING_XML)
        archive.writestr("word/_rels/document.xml.rels", _DOCUMENT_RELS_XML)
    return output.getvalue()


def re_numbered(line):
    head, separator, _ = line.partition(". ")
    return bool(separator and head.isdigit())


def _ooxml_paragraph(text, style="Normal", numbering=None):
    props = [f'<w:pStyle w:val="{escape(style)}"/>'] if style else []
    if numbering:
        num_id, level = numbering
        props.append(f'<w:numPr><w:ilvl w:val="{level}"/><w:numId w:val="{num_id}"/></w:numPr>')
    ppr = f"<w:pPr>{''.join(props)}</w:pPr>" if props else ""
    return f'<w:p>{ppr}<w:r><w:t xml:space="preserve">{escape(str(text or ""))}</w:t></w:r></w:p>'


_CONTENT_TYPES_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
<Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>
<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
</Types>'''

_ROOT_RELS_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>'''

_DOCUMENT_RELS_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering" Target="numbering.xml"/>
</Relationships>'''

_CORE_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"><dc:creator></dc:creator><cp:lastModifiedBy></cp:lastModifiedBy></cp:coreProperties>'''

_APP_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"><Application>MMN BF Factory</Application></Properties>'''

_STYLES_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:docDefaults><w:rPrDefault><w:rPr><w:rFonts w:ascii="{{DOCX_FONT}}" w:hAnsi="{{DOCX_FONT}}" w:eastAsia="{{DOCX_FONT}}" w:cs="{{DOCX_FONT}}" w:hint="eastAsia"/><w:sz w:val="22"/><w:lang w:val="zh-CN" w:eastAsia="zh-CN"/></w:rPr></w:rPrDefault><w:pPrDefault><w:pPr><w:spacing w:after="120" w:line="300" w:lineRule="auto"/></w:pPr></w:pPrDefault></w:docDefaults>
<w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/></w:style>
<w:style w:type="paragraph" w:styleId="Kicker"><w:name w:val="Kicker"/><w:basedOn w:val="Normal"/><w:rPr><w:color w:val="145B6E"/><w:b/><w:sz w:val="18"/></w:rPr></w:style>
<w:style w:type="paragraph" w:styleId="DocTitle"><w:name w:val="Document Title"/><w:basedOn w:val="Normal"/><w:pPr><w:spacing w:after="120"/></w:pPr><w:rPr><w:b/><w:sz w:val="56"/><w:color w:val="1A2227"/></w:rPr></w:style>
<w:style w:type="paragraph" w:styleId="Subtitle"><w:name w:val="Subtitle"/><w:basedOn w:val="Normal"/><w:pPr><w:spacing w:after="360"/></w:pPr><w:rPr><w:sz w:val="24"/><w:color w:val="5A676E"/></w:rPr></w:style>
<w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:pPr><w:keepNext/><w:spacing w:before="360" w:after="200"/></w:pPr><w:rPr><w:b/><w:sz w:val="32"/><w:color w:val="145B6E"/></w:rPr></w:style>
<w:style w:type="paragraph" w:styleId="Quote"><w:name w:val="Quote"/><w:basedOn w:val="Normal"/><w:pPr><w:ind w:left="360"/></w:pPr><w:rPr><w:i/><w:color w:val="5A676E"/></w:rPr></w:style>
</w:styles>'''

_NUMBERING_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:abstractNum w:abstractNumId="1"><w:lvl w:ilvl="0"><w:numFmt w:val="bullet"/><w:lvlText w:val="•"/><w:pPr><w:tabs><w:tab w:val="num" w:pos="270"/></w:tabs><w:ind w:left="540" w:hanging="270"/></w:pPr></w:lvl></w:abstractNum>
<w:abstractNum w:abstractNumId="2"><w:lvl w:ilvl="0"><w:numFmt w:val="decimal"/><w:lvlText w:val="%1."/><w:pPr><w:tabs><w:tab w:val="num" w:pos="270"/></w:tabs><w:ind w:left="540" w:hanging="270"/></w:pPr></w:lvl></w:abstractNum>
<w:num w:numId="1"><w:abstractNumId w:val="1"/></w:num><w:num w:numId="2"><w:abstractNumId w:val="2"/></w:num>
</w:numbering>'''
